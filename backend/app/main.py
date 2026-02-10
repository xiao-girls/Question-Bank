from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Optional, List
import uvicorn
import logging
import tempfile
import os
from docx import Document
from docx.shared import Inches, Pt

from app.models.model_client import ModelClient, ModelConfig
from app.knowledge.knowledge_base import KnowledgeManager
from app.generator.question_generator import QuestionGenerator


# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)

# 设置控制台输出编码
import sys
sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')
logger = logging.getLogger(__name__)


app = FastAPI(title="题库生成系统 API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

knowledge_manager = KnowledgeManager()
model_client: Optional[ModelClient] = None
question_generator: Optional[QuestionGenerator] = None


class ModelConfigRequest(BaseModel):
    model_type: str
    base_url: str
    api_key: Optional[str] = None
    model_name: str


class KnowledgeBaseCreate(BaseModel):
    name: str
    description: str


class QuestionGenerateRequest(BaseModel):
    kb_id: str
    keyword: str
    question_type: str
    difficulty: str
    count: int = 5


@app.on_event("startup")
async def startup_event():
    global model_client, question_generator
    try:
        default_config = ModelConfig(
            model_type="ollama",
            base_url="http://localhost:11434",
            model_name="qwen3:1.7b"
        )
        model_client = ModelClient(default_config)
        question_generator = QuestionGenerator(model_client)
        logger.info("服务启动成功，模型客户端已初始化")
    except Exception as e:
        logger.error(f"服务启动失败: {str(e)}", exc_info=True)


@app.get("/")
async def root():
    return {"message": "题库生成系统 API", "version": "1.0.0"}


@app.get("/health")
async def health_check():
    return {"status": "healthy"}


@app.post("/api/model/config")
async def set_model_config(config: ModelConfigRequest):
    global model_client, question_generator
    try:
        await model_client.close()
        
        model_config = ModelConfig(**config.model_dump())
        model_client = ModelClient(model_config)
        question_generator = QuestionGenerator(model_client)
        
        logger.info(f"模型配置已更新: {config.model_type} - {config.model_name}")
        return {"message": "模型配置已更新", "config": config.model_dump()}
    except Exception as e:
        logger.error(f"更新模型配置失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"更新模型配置失败: {str(e)}")


@app.get("/api/model/config")
async def get_model_config():
    try:
        return {
            "model_type": model_client.config.model_type,
            "base_url": model_client.config.base_url,
            "model_name": model_client.config.model_name
        }
    except Exception as e:
        logger.error(f"获取模型配置失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"获取模型配置失败: {str(e)}")


@app.post("/api/knowledge")
async def create_knowledge_base(kb: KnowledgeBaseCreate):
    try:
        kb_id = knowledge_manager.create_knowledge_base(kb.name, kb.description)
        logger.info(f"创建知识库成功: {kb_id} - {kb.name}")
        return {"kb_id": kb_id, "message": "知识库创建成功"}
    except Exception as e:
        logger.error(f"创建知识库失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"创建知识库失败: {str(e)}")


@app.get("/api/knowledge")
async def list_knowledge_bases():
    try:
        return {"knowledge_bases": knowledge_manager.list_knowledge_bases()}
    except Exception as e:
        logger.error(f"获取知识库列表失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"获取知识库列表失败: {str(e)}")


@app.get("/api/knowledge/{kb_id}")
async def get_knowledge_base(kb_id: str):
    try:
        kb = knowledge_manager.get_knowledge_base(kb_id)
        if not kb:
            raise HTTPException(status_code=404, detail="知识库不存在")
        return kb
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取知识库失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"获取知识库失败: {str(e)}")


@app.post("/api/knowledge/{kb_id}/upload")
async def upload_file(kb_id: str, file: UploadFile = File(...)):
    try:
        content = await file.read()
        text_content = content.decode("utf-8", errors="ignore")
        
        success = knowledge_manager.upload_file(kb_id, file.filename, text_content)
        if not success:
            raise HTTPException(status_code=404, detail="知识库不存在")
        
        logger.info(f"文件上传成功: {kb_id} - {file.filename}")
        return {"message": "文件上传成功", "filename": file.filename}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"上传文件失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"上传文件失败: {str(e)}")


@app.delete("/api/knowledge/{kb_id}")
async def delete_knowledge_base(kb_id: str):
    try:
        success = knowledge_manager.delete_knowledge_base(kb_id)
        if not success:
            raise HTTPException(status_code=404, detail="知识库不存在")
        logger.info(f"知识库删除成功: {kb_id}")
        return {"message": "知识库删除成功"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"删除知识库失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"删除知识库失败: {str(e)}")


@app.put("/api/knowledge/{kb_id}")
async def update_knowledge_base(kb_id: str, kb: KnowledgeBaseCreate):
    try:
        success = knowledge_manager.update_knowledge_base(kb_id, kb.name, kb.description)
        if not success:
            raise HTTPException(status_code=404, detail="知识库不存在")
        logger.info(f"知识库更新成功: {kb_id} - {kb.name}")
        return {"message": "知识库更新成功", "name": kb.name}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"更新知识库失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"更新知识库失败: {str(e)}")


@app.delete("/api/knowledge/{kb_id}/files/{filename}")
async def delete_file(kb_id: str, filename: str):
    try:
        success = knowledge_manager.delete_file(kb_id, filename)
        if not success:
            raise HTTPException(status_code=404, detail="知识库或文件不存在")
        logger.info(f"文件删除成功: {kb_id} - {filename}")
        return {"message": "文件删除成功", "filename": filename}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"删除文件失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"删除文件失败: {str(e)}")


@app.post("/api/questions/generate")
async def generate_questions(request: QuestionGenerateRequest):
    try:
        logger.info(f"接收到题目生成请求: {request.kb_id} - {request.keyword}")
        
        if question_generator is None:
            raise HTTPException(status_code=500, detail="题目生成器未初始化")
        
        kb = knowledge_manager.get_knowledge_base(request.kb_id)
        if not kb:
            raise HTTPException(status_code=404, detail="知识库不存在")
        
        # 检查知识库是否有文件
        if not kb.get("files") or len(kb.get("files")) == 0:
            raise HTTPException(status_code=400, detail="知识库中没有文件，请先上传文件")
        
        knowledge_content = knowledge_manager.get_knowledge_content(request.kb_id)
        if not knowledge_content:
            raise HTTPException(status_code=400, detail="知识库内容为空")
        
        questions = await question_generator.generate_questions(
            knowledge_content=knowledge_content,
            keyword=request.keyword,
            question_type=request.question_type,
            difficulty=request.difficulty,
            count=request.count
        )
        
        if not questions:
            logger.warning(f"未生成题目: {request.keyword}")
            return {"questions": [], "message": "未生成题目，请检查配置和知识库内容"}
        
        logger.info(f"题目生成成功: {len(questions)}道")
        return {"questions": [q.model_dump() for q in questions]}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"生成题目失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"生成题目失败: {str(e)}")


class ExportRequest(BaseModel):
    questions: List[dict]
    title: Optional[str] = "题库导出"


@app.post("/api/questions/export")
async def export_questions(request: ExportRequest):
    try:
        logger.info(f"接收到题目导出请求: {len(request.questions)}道题目")
        
        # 创建一个临时文件
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
            temp_file_path = temp_file.name
        
        # 创建 Word 文档
        doc = Document()
        
        # 设置文档默认字体为微软雅黑
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(12)
        
        # 设置标题字体
        heading_style = doc.styles['Heading 1']
        heading_font = heading_style.font
        heading_font.name = 'Microsoft YaHei'
        heading_font.size = Pt(16)
        heading_font.bold = True
        
        heading2_style = doc.styles['Heading 2']
        heading2_font = heading2_style.font
        heading2_font.name = 'Microsoft YaHei'
        heading2_font.size = Pt(14)
        heading2_font.bold = True
        
        # 添加标题
        doc.add_heading(request.title, level=0)
        doc.add_paragraph()
        
        # 题目类型和难度的中文映射
        type_mapping = {
            "single_choice": "单选题",
            "multiple_choice": "多选题",
            "fill": "填空题",
            "essay": "问答题",
            "true_false": "判断题"
        }
        
        difficulty_mapping = {
            "easy": "简单",
            "medium": "中等",
            "hard": "困难"
        }
        
        # 添加题目
        for i, question in enumerate(request.questions, 1):
            # 题目编号和类型
            q_type = type_mapping.get(question.get("type"), question.get("type"))
            q_difficulty = difficulty_mapping.get(question.get("difficulty"), question.get("difficulty"))
            
            # 添加题目标题
            doc.add_heading(f"第{i}题 [{q_type}] [{q_difficulty}]", level=2)
            
            # 添加题目内容
            doc.add_paragraph(question.get("question"))
            
            # 添加选项（如果有）
            options = question.get("options")
            if options:
                for j, option in enumerate(options):
                    doc.add_paragraph(f"{chr(65 + j)}. {option}", style="List Bullet")
            
            # 添加答案
            doc.add_paragraph(f"答案: {question.get('answer')}")
            
            # 添加解析（如果有）
            explanation = question.get("explanation")
            if explanation:
                doc.add_paragraph(f"解析: {explanation}")
            
            # 添加空行
            doc.add_paragraph()
        
        # 保存文档
        doc.save(temp_file_path)
        logger.info(f"Word 文档生成成功: {temp_file_path}")
        
        # 返回文件
        return FileResponse(
            path=temp_file_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{request.title}.docx",
            background=True
        )
    except Exception as e:
        logger.error(f"导出题目失败: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"导出题目失败: {str(e)}")


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
